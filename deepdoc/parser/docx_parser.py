#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from docx import Document
import re
import logging
import pandas as pd
from collections import Counter
from rag.nlp import rag_tokenizer
from io import BytesIO
from common.text_quality import is_gibberish, filter_gibberish


class RAGFlowDocxParser:

    def __extract_table_content(self, tb):
        # 使用列表推导式，提升性能
        df = [[c.text for c in row.cells] for row in tb.rows]
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):

        def blockType(b):
            pattern = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in pattern:
                if re.search(p, b):
                    return n
            tks = [t for t in rag_tokenizer.tokenize(b).split() if len(t) > 1]
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"

            if len(tks) == 1 and rag_tokenizer.tag(tks[0]) == "nr":
                return "Nr"

            return "Ot"

        if len(df) < 2:
            return []
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        colnm = len(df.iloc[0, :])
        hdrows = [0]  # header is not necessarily appear in the first line
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)

        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))

        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000000):
        self.doc = Document(fnm) if isinstance(
            fnm, str) else Document(BytesIO(fnm))
        pn = 0 # parsed page
        secs = [] # parsed contents
        for p in self.doc.paragraphs:
            if pn > to_page:
                break

            runs_within_single_paragraph = [] # save runs within the range of pages
            for run in p.runs:
                if pn > to_page:
                    break
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text) # append run.text first

                # wrap page break checker into a static method
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1

            text = "".join(runs_within_single_paragraph)
            style = p.style.name if hasattr(p.style, 'name') else ''
            
            # Filter gibberish content
            if text.strip():
                if not is_gibberish(text):
                    filtered_text = filter_gibberish(text)
                    if filtered_text:
                        secs.append((filtered_text, style))
                else:
                    logging.debug(f"Filtered gibberish paragraph from DOCX: {text[:100]}...")
            else:
                secs.append((text, style))

        tbls = []
        for tb in self.doc.tables:
            table_content = self.__extract_table_content(tb)
            if table_content:
                # Filter gibberish in table content
                filtered_table_content = []
                for line in table_content:
                    if line.strip():
                        if not is_gibberish(line):
                            filtered_line = filter_gibberish(line)
                            if filtered_line:
                                filtered_table_content.append(filtered_line)
                        else:
                            logging.debug(f"Filtered gibberish table line from DOCX: {line[:100]}...")
                if filtered_table_content:
                    tbls.append(filtered_table_content)
        return secs, tbls
